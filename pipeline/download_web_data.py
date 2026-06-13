import os
import io
import re
import json
import time
import uuid
import hashlib
from pathlib import Path
from urllib.parse import urljoin, urlparse, urlunparse
from collections import deque

import requests
import numpy as np
import pandas as pd
from bs4 import BeautifulSoup, ParserRejectedMarkup
from tqdm import tqdm
from PIL import Image

import fitz  # PyMuPDF
import pytesseract
from docx import Document

from sqlalchemy import create_engine, text, inspect
from sqlalchemy.engine import URL

from qdrant_client import QdrantClient
from qdrant_client.http import models
from sentence_transformers import SentenceTransformer


# ============================================================
# KONFIGURACJA
# ============================================================

BASE = "https://bip.lublin.eu"

START_URLS = [
    "https://bip.lublin.eu/e-urzad/opisy-uslug/",
    "https://bip.lublin.eu/e-urzad/opisy-uslug/wyszukiwarka-uslug/",
    "https://bip.lublin.eu/prezydent-zastepcy-pelnomocnicy/prezydent/krzysztof-zuk-prezydent-miasta-lublin,1,14981,2.html",
    "https://bip.lublin.eu/prezydent-zastepcy-pelnomocnicy/zastepcy-prezydenta/",
    "https://bip.lublin.eu/prezydent-zastepcy-pelnomocnicy/pelnomocnicy/pelnomocnicy-prezydenta-miasta-lublin,1,14983,2.html",
    "https://bip.lublin.eu/prezydent-zastepcy-pelnomocnicy/organy-doradcze/",
    "https://bip.lublin.eu/urzad-miasta-lublin/struktura-organizacyjna/",
]

ALLOWED_PATH_PREFIXES = [
    "/e-urzad/opisy-uslug/",
    "/prezydent-zastepcy-pelnomocnicy/",
    "/urzad-miasta-lublin/struktura-organizacyjna/",
]
def env_int(name, default=None):
    value = os.getenv(name)
    if value is None or str(value).strip() == "":
        return default
    return int(value)


def env_float(name, default):
    value = os.getenv(name)
    if value is None or str(value).strip() == "":
        return default
    return float(value)


def env_bool(name, default=False):
    value = os.getenv(name)
    if value is None or str(value).strip() == "":
        return default
    return str(value).lower() in ["true", "1", "yes", "y"]


MAX_PAGES = env_int("PARSER_MAX_PAGES", None)
SLEEP_SECONDS = env_float("PARSER_SLEEP_SECONDS", 0.3)
REQUEST_TIMEOUT = env_int("PARSER_REQUEST_TIMEOUT", 30)

REPROCESS_ALL_DOCUMENTS = env_bool("PARSER_REPROCESS_ALL_DOCUMENTS", False)
MAX_DOCUMENTS_TO_PROCESS = env_int("PARSER_MAX_DOCUMENTS_TO_PROCESS", None)

DOWNLOAD_DIR = Path(os.getenv("PARSER_DOWNLOAD_DIR", "/data/bip_downloaded_documents"))
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)


# PostgreSQL - używamy istniejących zmiennych z .env
POSTGRES_HOST = os.getenv("DB_HOST", "postgres")
POSTGRES_PORT = int(os.getenv("DB_PORT", "5432"))
POSTGRES_DB = os.getenv("DB_NAME", "hackathon")
POSTGRES_USER = os.getenv("DB_USER", "postgres")
POSTGRES_PASSWORD = os.getenv("DB_PASSWORD", "postgres")


# Qdrant - używamy istniejących zmiennych z .env
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")
QDRANT_COLLECTION = os.getenv("QDRANT_COLLECTION", "urbanlab")
VECTOR_SIZE = int(os.getenv("EMBEDDING_DIMENSION", "384"))


# Model embeddingów - używamy istniejącej zmiennej
MODEL_NAME = os.getenv(
    "EMBEDDING_MODEL",
    "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
)


# OCR - używamy istniejącej zmiennej OCR_LANGUAGE
OCR_LANG = os.getenv("OCR_LANGUAGE", "pol+eng")
OCR_ZOOM = env_float("PARSER_OCR_ZOOM", 2.0)

TESSERACT_PATH = os.getenv("PARSER_TESSERACT_PATH", "/usr/bin/tesseract")
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

DOCUMENT_EXTENSIONS = {
    "pdf", "doc", "docx", "xls", "xlsx", "odt", "ods", "rtf",
    "txt", "csv", "zip", "7z", "rar"
}


HEADERS = {
    "User-Agent": "Mozilla/5.0",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "pl-PL,pl;q=0.9,en;q=0.8",
}


# ============================================================
# POŁĄCZENIA
# ============================================================

connection_url = URL.create(
    drivername="postgresql+psycopg2",
    username=POSTGRES_USER,
    password=POSTGRES_PASSWORD,
    host=POSTGRES_HOST,
    port=POSTGRES_PORT,
    database=POSTGRES_DB,
)

engine = create_engine(connection_url)
qdrant = QdrantClient(url=QDRANT_URL)

model = SentenceTransformer(MODEL_NAME)


# ============================================================
# FUNKCJE OGÓLNE
# ============================================================

def make_id(value: str) -> str:
    return hashlib.sha256(str(value).encode("utf-8")).hexdigest()


def qdrant_point_id(chunk_id: str) -> str:
    return str(uuid.uuid5(uuid.NAMESPACE_URL, str(chunk_id)))


def clean_text(text: str) -> str:
    if text is None:
        return ""
    text = str(text)
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_url(url: str) -> str:
    url = urljoin(BASE, str(url))
    parsed = urlparse(url)

    parsed = parsed._replace(fragment="")
    path = re.sub(r"/{2,}", "/", parsed.path)

    return urlunparse((
        parsed.scheme,
        parsed.netloc,
        path,
        "",
        parsed.query,
        ""
    ))


def get_extension_from_url(url: str) -> str:
    path = urlparse(str(url)).path
    return Path(path).suffix.lower().replace(".", "")


def is_document_url(url: str) -> bool:
    ext = get_extension_from_url(url)
    return ext in DOCUMENT_EXTENSIONS


def should_visit(url: str) -> bool:
    parsed = urlparse(url)

    if parsed.netloc and parsed.netloc != "bip.lublin.eu":
        return False

    path = parsed.path

    if is_document_url(url):
        return False

    skip_fragments = [
        "/wyszukiwarka/",
        "/search/",
        "/drukuj",
        "/print",
        "/wersja-do-druku",
        "/historia",
        "/zmiany",
        "/rejestr-zmian",
    ]

    if any(fragment in path.lower() for fragment in skip_fragments):
        return False

    return any(path.startswith(prefix) for prefix in ALLOWED_PATH_PREFIXES)


def fetch_html(url: str):
    try:
        response = requests.get(
            url,
            headers=HEADERS,
            timeout=REQUEST_TIMEOUT,
            allow_redirects=True
        )
        response.raise_for_status()

        content_type = response.headers.get("Content-Type", "").lower()
        if "text/html" not in content_type and "application/xhtml" not in content_type:
            return None

        return response.text

    except Exception:
        return None


def get_title(soup: BeautifulSoup) -> str:
    h1 = soup.find("h1")
    if h1:
        return clean_text(h1.get_text(" ", strip=True))

    title = soup.find("title")
    if title:
        return clean_text(title.get_text(" ", strip=True))

    return ""


def extract_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    main = soup.find("main") or soup.find("article") or soup.find("body") or soup

    lines = []
    for element in main.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th", "div"]):
        txt = clean_text(element.get_text(" ", strip=True))
        if txt and len(txt) > 2:
            lines.append(txt)

    result = "\n".join(lines)
    result = re.sub(r"\n{3,}", "\n\n", result)

    return clean_text(result)


def extract_links(soup: BeautifulSoup, source_url: str):
    links = []

    for a in soup.find_all("a", href=True):
        href = a.get("href")
        text_value = clean_text(a.get_text(" ", strip=True))
        target_url = normalize_url(urljoin(source_url, href))
        ext = get_extension_from_url(target_url)

        links.append({
            "text": text_value,
            "url": target_url,
            "is_document": ext in DOCUMENT_EXTENSIONS,
            "file_extension": ext,
        })

    return links


def classify_page_type(url: str, title: str, text_value: str) -> str:
    path = urlparse(url).path.lower()

    if "/wyszukiwarka-uslug/" in path:
        if re.search(r",\d+,\d+,\d+\.html$", path):
            return "service_card"
        return "services_search"

    if "/e-urzad/opisy-uslug/" in path:
        if re.search(r",\d+,\d+,\d+\.html$", path):
            return "service_card"
        return "services_index"

    if "/prezydent/" in path:
        return "president"

    if "/zastepcy-prezydenta" in path:
        return "deputy_mayors"

    if "/pelnomocnicy" in path:
        return "plenipotentiaries"

    if "/organy-doradcze" in path:
        return "advisory_bodies"

    if "/struktura-organizacyjna" in path:
        return "organizational_structure"

    return "other"


def classify_document_type(link_text: str, url: str) -> str:
    text_lower = f"{link_text} {url}".lower()

    if "wniosek" in text_lower:
        return "application_form"

    if "formularz" in text_lower:
        return "form"

    if "załącznik" in text_lower or "zalacznik" in text_lower:
        return "attachment"

    return "document"


# ============================================================
# CRAWLER STRON
# ============================================================

def crawl_pages() -> pd.DataFrame:
    visited = set()
    queue = deque([normalize_url(u) for u in START_URLS])
    rows = []

    discovered = set(queue)

    while queue and len(visited) < MAX_PAGES:
        url = queue.popleft()

        if url in visited:
            continue

        if not should_visit(url):
            continue

        visited.add(url)

        print(f"[{len(visited)}/{len(discovered)}] Pobieram: {url}")

        html = fetch_html(url)

        if html is None:
            continue

        try:
            soup = BeautifulSoup(html, "html.parser")
        except ParserRejectedMarkup:
            continue

        title = get_title(soup)
        text_value = extract_text(soup)
        links = extract_links(soup, url)
        page_type = classify_page_type(url, title, text_value)

        rows.append({
            "url": url,
            "title": title,
            "page_type": page_type,
            "text": text_value,
            "links": links,
            "links_count": len(links),
            "documents_count": sum(1 for x in links if x["is_document"]),
            "content_hash": make_id(title + "\n" + text_value + "\n" + json.dumps(links, ensure_ascii=False)),
            "updated_at": pd.Timestamp.utcnow().isoformat(),
        })

        new_count = 0

        for link in links:
            target = link["url"]

            if link["is_document"]:
                continue

            if target not in discovered and should_visit(target):
                discovered.add(target)
                queue.append(target)
                new_count += 1

        print(f"OK | nowe URL-e: {new_count} | kolejka: {len(queue)}")

        time.sleep(SLEEP_SECONDS)

    return pd.DataFrame(rows)


# ============================================================
# BUDOWANIE DANYCH GRAFOWYCH
# ============================================================

def build_documents_df(pages_df: pd.DataFrame) -> pd.DataFrame:
    rows = []

    for _, page in pages_df.iterrows():
        source_url = page["url"]

        for link in page["links"]:
            if not link["is_document"]:
                continue

            document_url = link["url"]
            link_text = link["text"]
            ext = link["file_extension"]

            rows.append({
                "document_id": make_id(document_url),
                "document_url": document_url,
                "source_page_url": source_url,
                "title": link_text,
                "document_type": classify_document_type(link_text, document_url),
                "file_extension": ext,
            })

    df = pd.DataFrame(rows)

    if len(df) == 0:
        return pd.DataFrame(columns=[
            "document_id", "document_url", "source_page_url",
            "title", "document_type", "file_extension"
        ])

    return df.drop_duplicates(subset=["document_url"]).reset_index(drop=True)


def build_nodes_df(pages_df: pd.DataFrame, documents_df: pd.DataFrame) -> pd.DataFrame:
    rows = []

    for _, row in pages_df.iterrows():
        rows.append({
            "node_id": make_id(row["url"]),
            "node_type": "Page",
            "url": row["url"],
            "title": row["title"],
            "text": row["text"],
            "page_type": row["page_type"],
            "content_hash": row["content_hash"],
        })

    for _, row in documents_df.iterrows():
        rows.append({
            "node_id": row["document_id"],
            "node_type": "Document",
            "url": row["document_url"],
            "title": row["title"],
            "text": "",
            "page_type": "",
            "content_hash": make_id(row["document_url"]),
        })

    return pd.DataFrame(rows)


def build_edges_df(pages_df: pd.DataFrame) -> pd.DataFrame:
    rows = []

    known_page_urls = set(pages_df["url"].tolist())

    for _, page in pages_df.iterrows():
        source_url = page["url"]
        source_id = make_id(source_url)

        for link in page["links"]:
            target_url = link["url"]

            if link["is_document"]:
                target_id = make_id(target_url)
                relation = "HAS_DOCUMENT"

                if classify_document_type(link["text"], target_url) == "application_form":
                    relation = "HAS_APPLICATION_FORM"

            else:
                if target_url not in known_page_urls:
                    continue

                target_id = make_id(target_url)
                relation = "LINKS_TO"

            evidence = link["text"]

            edge_id = make_id(f"{source_id}|{relation}|{target_id}|{evidence}")

            rows.append({
                "edge_id": edge_id,
                "source_id": source_id,
                "target_id": target_id,
                "relation": relation,
                "source_url": source_url,
                "target_url": target_url,
                "evidence": evidence,
            })

    df = pd.DataFrame(rows)

    if len(df) == 0:
        return pd.DataFrame(columns=[
            "edge_id", "source_id", "target_id", "relation",
            "source_url", "target_url", "evidence"
        ])

    return df.drop_duplicates(subset=["edge_id"]).reset_index(drop=True)


def split_long_paragraph(paragraph: str, max_chars: int = 2000):
    paragraph = clean_text(paragraph)

    if len(paragraph) <= max_chars:
        return [paragraph]

    sentences = re.split(r"(?<=[.!?])\s+", paragraph)
    chunks = []
    current = ""

    for sentence in sentences:
        if len(current) + len(sentence) + 1 <= max_chars:
            current = (current + " " + sentence).strip()
        else:
            if current:
                chunks.append(current)
            current = sentence

    if current:
        chunks.append(current)

    return chunks


def split_text_into_chunks(text_value: str, min_chars: int = 40, max_chars: int = 2000):
    text_value = clean_text(text_value)
    paragraphs = re.split(r"\n+", text_value)

    chunks = []

    for paragraph in paragraphs:
        paragraph = clean_text(paragraph)

        if len(paragraph) < min_chars:
            continue

        chunks.extend(split_long_paragraph(paragraph, max_chars=max_chars))

    return chunks


def build_page_chunks_df(nodes_df: pd.DataFrame) -> pd.DataFrame:
    rows = []

    pages = nodes_df[nodes_df["node_type"] == "Page"].copy()

    for _, row in pages.iterrows():
        chunks = split_text_into_chunks(row["text"])

        for chunk_index, chunk_text in enumerate(chunks):
            chunk_id = make_id(f'{row["node_id"]}:{chunk_index}:{chunk_text}')

            rows.append({
                "chunk_id": chunk_id,
                "node_id": row["node_id"],
                "chunk_index": chunk_index,
                "url": row["url"],
                "title": row["title"],
                "node_type": "Page",
                "source_type": "html_page",
                "text": chunk_text,
                "text_length": len(chunk_text),
                "source": "bip_lublin",
            })

    return pd.DataFrame(rows)


# ============================================================
# POSTGRES HELPERS
# ============================================================

def replace_table(df: pd.DataFrame, table_name: str, chunksize: int = 1000):
    df = df.where(pd.notnull(df), None)

    df.to_sql(
        table_name,
        engine,
        if_exists="replace",
        index=False,
        chunksize=chunksize
    )

    print(f"Zapisano tabelę: {table_name} | rows={len(df)}")


def table_exists(table_name: str) -> bool:
    return inspect(engine).has_table(table_name)


def delete_by_values(table_name: str, column_name: str, values, batch_size: int = 500):
    values = [str(v) for v in values if v is not None and str(v) != "nan"]
    values = list(dict.fromkeys(values))

    if not values:
        return

    with engine.begin() as conn:
        for start in range(0, len(values), batch_size):
            batch = values[start:start + batch_size]

            params = {f"v{i}": value for i, value in enumerate(batch)}
            placeholders = ", ".join([f":v{i}" for i in range(len(batch))])

            sql = text(f'DELETE FROM "{table_name}" WHERE "{column_name}" IN ({placeholders})')
            conn.execute(sql, params)


def append_table(df: pd.DataFrame, table_name: str, chunksize: int = 500):
    if len(df) == 0:
        return

    df = df.where(pd.notnull(df), None)

    df.to_sql(
        table_name,
        engine,
        if_exists="append",
        index=False,
        chunksize=chunksize
    )

    print(f"Dopisano do tabeli: {table_name} | rows={len(df)}")


def create_indexes():
    with engine.begin() as conn:
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_pages_url ON bip_pages(url);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_pages_page_type ON bip_pages(page_type);"))

        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_documents_url ON bip_documents(document_url);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_documents_source_page_url ON bip_documents(source_page_url);"))

        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_nodes_node_id ON bip_nodes(node_id);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_nodes_url ON bip_nodes(url);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_nodes_node_type ON bip_nodes(node_type);"))

        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_edges_source_id ON bip_edges(source_id);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_edges_target_id ON bip_edges(target_id);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_edges_relation ON bip_edges(relation);"))

        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_chunks_chunk_id ON bip_chunks(chunk_id);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_chunks_node_id ON bip_chunks(node_id);"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_chunks_url ON bip_chunks(url);"))

        if table_exists("bip_document_texts"):
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_texts_document_id ON bip_document_texts(document_id);"))
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_texts_document_url ON bip_document_texts(document_url);"))
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_texts_status ON bip_document_texts(status);"))

        if table_exists("bip_document_chunks"):
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_chunks_chunk_id ON bip_document_chunks(chunk_id);"))
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_chunks_document_id ON bip_document_chunks(document_id);"))
            conn.execute(text("CREATE INDEX IF NOT EXISTS idx_bip_document_chunks_url ON bip_document_chunks(url);"))

    print("Indeksy utworzone.")


# ============================================================
# QDRANT HELPERS
# ============================================================

def ensure_qdrant_collection():
    if not qdrant.collection_exists(QDRANT_COLLECTION):
        qdrant.create_collection(
            collection_name=QDRANT_COLLECTION,
            vectors_config=models.VectorParams(
                size=VECTOR_SIZE,
                distance=models.Distance.COSINE
            )
        )


def delete_qdrant_by_filter(filter_obj: models.Filter):
    qdrant.delete(
        collection_name=QDRANT_COLLECTION,
        points_selector=models.FilterSelector(
            filter=filter_obj
        ),
        wait=True
    )


def delete_html_pages_from_qdrant():
    """
    Usuwa stare embeddingi stron HTML.
    Łapie też stare punkty, gdzie source_type było None, ale node_type było Page.
    """
    ensure_qdrant_collection()

    filter_obj = models.Filter(
        must=[
            models.FieldCondition(
                key="node_type",
                match=models.MatchValue(value="Page")
            )
        ]
    )

    delete_qdrant_by_filter(filter_obj)
    print("Usunięto stare embeddingi stron HTML z Qdranta.")


def delete_document_from_qdrant(document_id: str):
    filter_obj = models.Filter(
        must=[
            models.FieldCondition(
                key="document_id",
                match=models.MatchValue(value=str(document_id))
            )
        ]
    )

    delete_qdrant_by_filter(filter_obj)


def upsert_chunks_to_qdrant(chunks_df: pd.DataFrame, batch_size: int = 256):
    if len(chunks_df) == 0:
        print("Brak chunków do Qdranta.")
        return

    ensure_qdrant_collection()

    embeddings = model.encode(
        chunks_df["text"].tolist(),
        batch_size=32,
        show_progress_bar=True,
        normalize_embeddings=True
    )

    embeddings = np.array(embeddings, dtype="float32")

    points = []

    records = chunks_df.to_dict(orient="records")

    for i, row in enumerate(records):
        payload = {
            "chunk_id": str(row.get("chunk_id", "")),
            "node_id": str(row.get("node_id", "")),
            "document_id": str(row.get("document_id", "")) if row.get("document_id") is not None else "",
            "chunk_index": int(row.get("chunk_index", 0)),
            "url": str(row.get("url", "")),
            "document_url": str(row.get("document_url", "")) if row.get("document_url") is not None else "",
            "source_page_url": str(row.get("source_page_url", "")) if row.get("source_page_url") is not None else "",
            "title": str(row.get("title", "")),
            "node_type": str(row.get("node_type", "")),
            "source_type": str(row.get("source_type", "")),
            "source": str(row.get("source", "bip_lublin")),
            "text": str(row.get("text", "")),
        }

        point = models.PointStruct(
            id=qdrant_point_id(row["chunk_id"]),
            vector=embeddings[i].tolist(),
            payload=payload
        )

        points.append(point)

        if len(points) >= batch_size:
            qdrant.upsert(
                collection_name=QDRANT_COLLECTION,
                points=points,
                wait=True
            )
            print(f"Wysłano do Qdranta: {i + 1}/{len(records)}")
            points = []

    if points:
        qdrant.upsert(
            collection_name=QDRANT_COLLECTION,
            points=points,
            wait=True
        )

    print("Zapis do Qdranta zakończony.")


# ============================================================
# OCR / EKSTRAKCJA DOKUMENTÓW
# ============================================================

def safe_filename_from_url(url: str, ext: str) -> str:
    file_hash = make_id(url)[:20]
    if ext:
        return f"{file_hash}.{ext}"
    return file_hash


def download_document(url: str, ext: str):
    filename = safe_filename_from_url(url, ext)
    path = DOWNLOAD_DIR / filename

    if path.exists() and path.stat().st_size > 0:
        return str(path), "cached", None

    try:
        response = requests.get(
            url,
            headers={"User-Agent": "Mozilla/5.0", "Accept": "*/*"},
            timeout=REQUEST_TIMEOUT,
            allow_redirects=True
        )
        response.raise_for_status()

        with open(path, "wb") as f:
            f.write(response.content)

        return str(path), "downloaded", None

    except Exception as e:
        return None, "download_error", str(e)


def extract_text_from_pdf(path: str):
    try:
        doc = fitz.open(path)

        direct_pages = []
        for page in doc:
            direct_pages.append(page.get_text("text"))

        direct_text = clean_text("\n\n".join(direct_pages))

        if len(direct_text) >= 100:
            return direct_text, "pdf_text", None

        ocr_pages = []

        for page in doc:
            pix = page.get_pixmap(
                matrix=fitz.Matrix(OCR_ZOOM, OCR_ZOOM),
                alpha=False
            )

            img = Image.open(io.BytesIO(pix.tobytes("png")))
            page_text = pytesseract.image_to_string(img, lang=OCR_LANG)
            ocr_pages.append(page_text)

        ocr_text = clean_text("\n\n".join(ocr_pages))

        if len(ocr_text) > 0:
            return ocr_text, "pdf_ocr", None

        return "", "pdf_empty", None

    except Exception as e:
        return "", "pdf_error", str(e)


def extract_text_from_docx(path: str):
    try:
        doc = Document(path)
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join([c for c in cells if c])
                if line:
                    parts.append(line)

        return clean_text("\n".join(parts)), "docx_text", None

    except Exception as e:
        return "", "docx_error", str(e)


def extract_text_from_excel(path: str):
    try:
        parts = []
        excel = pd.ExcelFile(path)

        for sheet_name in excel.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name, dtype=str)
            df = df.fillna("")

            parts.append(f"ARKUSZ: {sheet_name}")
            parts.append(df.to_csv(index=False, sep="\t"))

        return clean_text("\n\n".join(parts)), "excel_text", None

    except Exception as e:
        return "", "excel_error", str(e)


def extract_text_from_txt(path: str):
    try:
        for enc in ["utf-8", "cp1250", "latin-1"]:
            try:
                with open(path, "r", encoding=enc) as f:
                    return clean_text(f.read()), "txt_text", None
            except UnicodeDecodeError:
                pass

        return "", "txt_error", "Nie udało się odczytać kodowania."

    except Exception as e:
        return "", "txt_error", str(e)


def extract_document_text(path: str, ext: str):
    ext = (ext or "").lower()

    if ext == "pdf":
        return extract_text_from_pdf(path)

    if ext == "docx":
        return extract_text_from_docx(path)

    if ext in ["xlsx", "xls"]:
        return extract_text_from_excel(path)

    if ext in ["txt", "csv"]:
        return extract_text_from_txt(path)

    if ext == "doc":
        return "", "unsupported_doc", "Stary format .doc wymaga konwersji np. przez LibreOffice."

    return "", "unsupported_extension", f"Nieobsługiwane rozszerzenie: {ext}"


def process_documents_for_ocr(documents_df: pd.DataFrame):
    if len(documents_df) == 0:
        print("Brak dokumentów.")
        return pd.DataFrame(), pd.DataFrame()

    documents_unique = documents_df.drop_duplicates(subset=["document_url"]).copy()

    already_processed_ids = set()

    if table_exists("bip_document_texts") and not REPROCESS_ALL_DOCUMENTS:
        old = pd.read_sql("SELECT document_id FROM bip_document_texts", engine)
        already_processed_ids = set(old["document_id"].astype(str).tolist())

    if not REPROCESS_ALL_DOCUMENTS:
        documents_unique = documents_unique[
            ~documents_unique["document_id"].astype(str).isin(already_processed_ids)
        ].copy()

    if MAX_DOCUMENTS_TO_PROCESS is not None:
        documents_unique = documents_unique.head(MAX_DOCUMENTS_TO_PROCESS)

    print("Dokumenty do OCR/ekstrakcji:", len(documents_unique))

    document_text_rows = []

    for _, row in tqdm(documents_unique.iterrows(), total=len(documents_unique)):
        document_id = str(row["document_id"])
        document_url = str(row["document_url"])
        title = str(row.get("title", ""))
        source_page_url = str(row.get("source_page_url", ""))
        document_type = str(row.get("document_type", "document"))
        ext = str(row.get("file_extension", "")).lower().replace(".", "")

        if not ext or ext == "nan":
            ext = get_extension_from_url(document_url)

        local_path, download_status, download_error = download_document(document_url, ext)

        text_content = ""
        extraction_method = ""
        extraction_error = None
        status = "ok"

        if local_path is None:
            status = "download_error"
            extraction_method = "none"
            extraction_error = download_error
        else:
            text_content, extraction_method, extraction_error = extract_document_text(local_path, ext)

            if extraction_error:
                status = "extraction_error"

            if len(clean_text(text_content)) == 0 and status == "ok":
                status = "empty_text"

        text_content = clean_text(text_content)

        document_text_rows.append({
            "document_id": document_id,
            "document_url": document_url,
            "source_page_url": source_page_url,
            "title": title,
            "document_type": document_type,
            "file_extension": ext,
            "local_path": local_path,
            "download_status": download_status,
            "extraction_method": extraction_method,
            "status": status,
            "error": extraction_error,
            "text": text_content,
            "text_length": len(text_content),
            "content_hash": make_id(text_content),
            "updated_at": pd.Timestamp.utcnow().isoformat(),
        })

        time.sleep(SLEEP_SECONDS)

    document_texts_df = pd.DataFrame(document_text_rows)

    if len(document_texts_df) == 0:
        return document_texts_df, pd.DataFrame()

    document_chunk_rows = []

    for _, row in document_texts_df.iterrows():
        if row["status"] not in ["ok", "empty_text"]:
            continue

        text_content = row["text"]

        if not text_content or len(text_content) < 40:
            continue

        chunks = split_text_into_chunks(text_content)

        for chunk_index, chunk_text in enumerate(chunks):
            chunk_id = make_id(f'{row["document_id"]}:{chunk_index}:{chunk_text}')

            document_chunk_rows.append({
                "chunk_id": chunk_id,
                "document_id": row["document_id"],
                "node_id": row["document_id"],
                "chunk_index": chunk_index,
                "document_url": row["document_url"],
                "url": row["document_url"],
                "source_page_url": row["source_page_url"],
                "title": row["title"],
                "document_type": row["document_type"],
                "file_extension": row["file_extension"],
                "node_type": "Document",
                "source_type": "document_ocr",
                "text": chunk_text,
                "text_length": len(chunk_text),
                "source": "bip_lublin",
            })

    document_chunks_df = pd.DataFrame(document_chunk_rows)

    return document_texts_df, document_chunks_df


# ============================================================
# GŁÓWNA AKTUALIZACJA
# ============================================================

def run_update():
    print("Test połączeń...")

    with engine.begin() as conn:
        print("Postgres:", conn.execute(text("SELECT current_database(), current_user;")).fetchone())

    print("Qdrant:", qdrant.get_collections())

    print("\n=== 1. Crawl stron BIP ===")
    pages_df = crawl_pages()

    print("pages_df:", pages_df.shape)

    print("\n=== 2. Budowa dokumentów, nodes, edges, chunks ===")
    documents_df = build_documents_df(pages_df)
    nodes_df = build_nodes_df(pages_df, documents_df)
    edges_df = build_edges_df(pages_df)
    page_chunks_df = build_page_chunks_df(nodes_df)

    print("documents_df:", documents_df.shape)
    print("nodes_df:", nodes_df.shape)
    print("edges_df:", edges_df.shape)
    print("page_chunks_df:", page_chunks_df.shape)

    print("\n=== 3. Aktualizacja PostgreSQL - strony/graf ===")
    pages_to_sql = pages_df.copy()
    pages_to_sql["links"] = pages_to_sql["links"].apply(
        lambda x: json.dumps(x, ensure_ascii=False) if isinstance(x, list) else "[]"
    )

    replace_table(pages_to_sql, "bip_pages", chunksize=500)
    replace_table(documents_df, "bip_documents", chunksize=500)
    replace_table(nodes_df, "bip_nodes", chunksize=500)
    replace_table(edges_df, "bip_edges", chunksize=1000)
    replace_table(page_chunks_df, "bip_chunks", chunksize=1000)

    print("\n=== 4. Aktualizacja Qdrant - strony HTML ===")
    delete_html_pages_from_qdrant()
    upsert_chunks_to_qdrant(page_chunks_df, batch_size=256)

    print("\n=== 5. OCR / ekstrakcja nowych dokumentów ===")
    document_texts_df, document_chunks_df = process_documents_for_ocr(documents_df)

    print("document_texts_df:", document_texts_df.shape)
    print("document_chunks_df:", document_chunks_df.shape)

    if len(document_texts_df) > 0:
        processed_ids = document_texts_df["document_id"].astype(str).tolist()

        print("\n=== 6. Aktualizacja PostgreSQL - dokumenty OCR ===")

        if table_exists("bip_document_texts"):
            delete_by_values("bip_document_texts", "document_id", processed_ids)

        if table_exists("bip_document_chunks"):
            delete_by_values("bip_document_chunks", "document_id", processed_ids)

        append_table(document_texts_df, "bip_document_texts", chunksize=100)

        if len(document_chunks_df) > 0:
            append_table(document_chunks_df, "bip_document_chunks", chunksize=500)

        print("\n=== 7. Aktualizacja Qdrant - dokumenty OCR ===")

        for document_id in processed_ids:
            delete_document_from_qdrant(document_id)

        if len(document_chunks_df) > 0:
            upsert_chunks_to_qdrant(document_chunks_df, batch_size=256)

    else:
        print("Brak nowych dokumentów do OCR/ekstrakcji.")

    print("\n=== 8. Indeksy ===")
    create_indexes()

    print("\n=== 9. Sprawdzenie ===")

    checks = [
        "bip_pages",
        "bip_documents",
        "bip_nodes",
        "bip_edges",
        "bip_chunks",
        "bip_document_texts",
        "bip_document_chunks",
    ]

    for table_name in checks:
        if table_exists(table_name):
            df_count = pd.read_sql(f"SELECT COUNT(*) AS count FROM {table_name}", engine)
            print(table_name, "=", int(df_count.iloc[0]["count"]))

    print("\nQdrant:")
    print(qdrant.get_collection(QDRANT_COLLECTION))

    print("\nAKTUALIZACJA ZAKOŃCZONA.")


# ============================================================
# URUCHOMIENIE
# ============================================================

if __name__ == "__main__":
    run_update()